"""Documents module service."""

from __future__ import annotations

import hashlib
import re
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Any, Literal
from uuid import UUID, uuid4

from psycopg.errors import UniqueViolation
from psycopg.types.json import Jsonb
from pydantic import BaseModel, Field

from app.db import transaction
from app.lifecycle import (
    LifeItemError,
    create_life_item,
    delete_life_item,
    process_lifecycle_for_item,
    set_lifecycle_status,
)
from app.llm import LLMUnavailable, generate_json
from app.rag import embed_documents
from app.rag.retrieval import vector_literal


class DocumentNameConflictError(ValueError):
    """Raised when a document unique_name already exists."""


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be converted to text."""


class DocumentCreate(BaseModel):
    original_name: str = Field(min_length=1)
    content: str = Field(min_length=1)
    unique_name: str | None = None
    mime_type: str = "text/plain"
    request_id: str | None = None
    source: dict[str, Any] = Field(default_factory=dict)


class DocumentRename(BaseModel):
    unique_name: str = Field(pattern=r"^[a-z][a-z0-9_]*$")


class DocumentAnnotation(BaseModel):
    category_tag: str | None = Field(default=None, max_length=40)
    connection_summary: str | None = Field(default=None, max_length=400)


class DocumentItem(BaseModel):
    id: UUID
    title: str
    description: str
    lifecycle_status: str
    connection_status: str
    chunk_status: str
    bucket_update_status: str
    unique_name: str
    original_name: str
    mime_type: str
    byte_size: int
    content_sha256: str
    category_tag: str
    connection_summary: str
    tag_status: str
    created_at: datetime
    updated_at: datetime


def create_document(
    payload: DocumentCreate,
    *,
    review: bool = True,
    review_root: Path | None = None,
) -> DocumentItem:
    content_sha256 = hashlib.sha256(payload.content.encode("utf-8")).hexdigest()
    byte_size = len(payload.content.encode("utf-8"))
    unique_name = payload.unique_name or _generate_available_unique_name(payload.original_name, payload.content)
    request_id = payload.request_id or f"document-{uuid4().hex}"
    summary = _summary(payload.content)
    category_tag, connection_summary, tag_status = _annotate_document(payload.content, summary=summary)

    try:
        result = create_life_item(
            module_id="documents",
            item_type="document",
            title=payload.original_name,
            description=summary,
            payload={
                "unique_name": unique_name,
                "original_name": payload.original_name,
                "mime_type": payload.mime_type,
                "byte_size": byte_size,
                "content_sha256": content_sha256,
                "summary": summary,
                "category_tag": category_tag,
                "connection_summary": connection_summary,
                "tag_status": tag_status,
            },
            source={
                "kind": "manual_document",
                **payload.source,
            },
            request_id=request_id,
            side_table_data={
                "unique_name": unique_name,
                "original_name": payload.original_name,
                "mime_type": payload.mime_type,
                "byte_size": byte_size,
                "content_sha256": content_sha256,
                "category_tag": category_tag,
                "connection_summary": connection_summary,
                "tag_status": tag_status,
            },
        )
    except UniqueViolation as exc:
        raise DocumentNameConflictError(unique_name) from exc

    if result.created:
        _create_document_chunks(result.item["id"], payload.content, unique_name)
        if review:
            process_lifecycle_for_item(result.item["id"], root=review_root)
            _set_document_bucket_update_text(result.item["id"], connection_summary)
            _auto_weave_connected_buckets(result.item["id"])

    return get_document(result.item["id"])


def create_document_from_upload(
    *,
    original_name: str,
    data: bytes,
    mime_type: str | None,
    review: bool = True,
    review_root: Path | None = None,
) -> DocumentItem:
    content = _extract_upload_text(original_name, data)
    request_id = f"document-upload-{hashlib.sha256(data).hexdigest()}"
    return create_document(
        DocumentCreate(
            original_name=original_name,
            content=content,
            mime_type=mime_type or _guess_mime_type(original_name),
            request_id=request_id,
            source={"kind": "file_upload", "original_name": original_name},
        ),
        review=review,
        review_root=review_root,
    )


def list_documents(
    *,
    status: Literal["active", "archived", "deleted"] | None = "active",
    limit: int = 50,
) -> list[DocumentItem]:
    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT li.*, di.unique_name, di.original_name, di.mime_type,
                    di.byte_size, di.content_sha256, di.category_tag,
                    di.connection_summary, di.tag_status
                FROM life_items li
                JOIN document_items di ON di.life_item_id = li.id
                JOIN module_instances mi ON mi.id = li.module_instance_id
                JOIN modules m ON m.id = mi.module_id
                WHERE m.id = 'documents'
                    AND (%(status)s::text IS NULL OR li.lifecycle_status = %(status)s)
                ORDER BY li.created_at DESC
                LIMIT %(limit)s
                """,
                {"status": status, "limit": limit},
            )
            return [_row_to_document(row) for row in cur.fetchall()]


def get_document(document_id: UUID | str) -> DocumentItem:
    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT li.*, di.unique_name, di.original_name, di.mime_type,
                    di.byte_size, di.content_sha256, di.category_tag,
                    di.connection_summary, di.tag_status
                FROM life_items li
                JOIN document_items di ON di.life_item_id = li.id
                JOIN module_instances mi ON mi.id = li.module_instance_id
                JOIN modules m ON m.id = mi.module_id
                WHERE li.id = %s AND m.id = 'documents'
                """,
                (document_id,),
            )
            row = cur.fetchone()
            if row is None:
                raise LifeItemError(f"Unknown Document: {document_id}")
            return _row_to_document(row)


def rename_document(document_id: UUID | str, unique_name: str) -> DocumentItem:
    document = get_document(document_id)
    try:
        with transaction() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    UPDATE document_items
                    SET unique_name = %s, updated_at = now()
                    WHERE life_item_id = %s
                    """,
                    (unique_name, document.id),
                )
                cur.execute(
                    """
                    UPDATE life_items
                    SET payload = jsonb_set(payload, '{unique_name}', to_jsonb(%s::text)),
                        updated_at = now()
                    WHERE id = %s
                    """,
                    (unique_name, document.id),
                )
                cur.execute(
                    """
                    UPDATE knowledge_chunks
                    SET metadata = jsonb_set(metadata, '{unique_name}', to_jsonb(%s::text)),
                        updated_at = now()
                    WHERE life_item_id = %s
                    """,
                    (unique_name, document.id),
                )
    except UniqueViolation as exc:
        raise DocumentNameConflictError(unique_name) from exc

    return get_document(document.id)


def update_document_annotation(document_id: UUID | str, payload: DocumentAnnotation) -> DocumentItem:
    document = get_document(document_id)
    category_tag = _slug_tag(payload.category_tag) if payload.category_tag is not None else None
    connection_summary = (
        _compact_text(payload.connection_summary, limit=400)
        if payload.connection_summary is not None
        else None
    )
    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                UPDATE document_items
                SET category_tag = COALESCE(%s, category_tag),
                    connection_summary = COALESCE(%s, connection_summary),
                    tag_status = 'complete',
                    updated_at = now()
                WHERE life_item_id = %s
                """,
                (category_tag, connection_summary, document.id),
            )
            cur.execute(
                """
                UPDATE life_items
                SET payload = payload
                        || jsonb_strip_nulls(jsonb_build_object(
                            'category_tag', %s::text,
                            'connection_summary', %s::text,
                            'tag_status', 'complete'
                        )),
                    updated_at = now()
                WHERE id = %s
                """,
                (category_tag, connection_summary, document.id),
            )

    return get_document(document.id)


def archive_document(document_id: UUID | str) -> DocumentItem:
    get_document(document_id)
    set_lifecycle_status(document_id, "archived")
    return get_document(document_id)


def remove_document(document_id: UUID | str) -> None:
    document = get_document(document_id)
    delete_life_item(document.id)


def _create_document_chunks(life_item_id: UUID, content: str, unique_name: str) -> None:
    chunks = _chunk_text(content)
    embeddings = embed_documents(chunks)
    with transaction() as conn:
        with conn.cursor() as cur:
            for index, chunk in enumerate(chunks):
                embedding = embeddings[index] if index < len(embeddings) else None
                metadata = {
                    "chunk_index": index,
                    "total_chunks": len(chunks),
                    "unique_name": unique_name,
                    "embedding_status": "complete" if embedding else "not_available",
                }
                if embedding:
                    cur.execute(
                        """
                        INSERT INTO knowledge_chunks (life_item_id, content, embedding, source_type, metadata)
                        VALUES (%s, %s, CAST(%s AS vector), 'document', %s)
                        """,
                        (
                            life_item_id,
                            chunk,
                            vector_literal(embedding),
                            Jsonb(metadata),
                        ),
                    )
                else:
                    cur.execute(
                        """
                        INSERT INTO knowledge_chunks (life_item_id, content, source_type, metadata)
                        VALUES (%s, %s, 'document', %s)
                        """,
                        (
                            life_item_id,
                            chunk,
                            Jsonb(metadata),
                        ),
                    )


def _row_to_document(row: dict[str, Any]) -> DocumentItem:
    return DocumentItem(
        id=row["id"],
        title=row["title"],
        description=row["description"],
        lifecycle_status=row["lifecycle_status"],
        connection_status=row["connection_status"],
        chunk_status=row["chunk_status"],
        bucket_update_status=row["bucket_update_status"],
        unique_name=row["unique_name"],
        original_name=row["original_name"],
        mime_type=row["mime_type"],
        byte_size=row["byte_size"],
        content_sha256=row["content_sha256"],
        category_tag=row["category_tag"],
        connection_summary=row["connection_summary"],
        tag_status=row["tag_status"],
        created_at=row["created_at"],
        updated_at=row["updated_at"],
    )


def _annotate_document(content: str, *, summary: str) -> tuple[str, str, str]:
    from app.user_model import list_goals, list_story_bucket_items

    context = {
        "story_buckets": [
            {"name": bucket.display_name, "description": bucket.description, "content": bucket.content[:500]}
            for bucket in list_story_bucket_items()
        ],
        "active_goals": [{"title": goal.title, "body": goal.body[:300]} for goal in list_goals() if goal.status == "active"],
    }
    try:
        response = generate_json(
            _annotation_prompt(content[:4000], context),
            system=(
                "You categorize a personal document and explain in one sentence how it connects "
                "to the person, using their story buckets and goals. Return only JSON: "
                '{"category_tag": str, "connection_summary": str}.'
            ),
            temperature=0.2,
            max_output_tokens=300,
        )
        tag = _slug_tag(str(response.get("category_tag") or ""))
        connection_summary = _compact_text(str(response.get("connection_summary") or ""), limit=280)
        if tag and connection_summary:
            return tag, connection_summary, "complete"
    except (LLMUnavailable, Exception):
        pass
    return _fallback_tag(content), _fallback_summary(summary), "failed"


def _annotation_prompt(content: str, context: dict[str, Any]) -> str:
    return (
        "Return JSON for this document annotation.\n"
        "Rules:\n"
        "- category_tag is lowercase, concise, and free-form.\n"
        "- connection_summary is one sentence under 280 characters.\n"
        "- Use the person context only when it is relevant.\n\n"
        f"document_text:\n{content}\n\n"
        f"user_context:\n{context}\n"
    )


def _slug_tag(value: str | None) -> str:
    compact = re.sub(r"[^a-z0-9]+", "-", (value or "").lower()).strip("-")
    compact = re.sub(r"-+", "-", compact)
    return compact[:40].strip("-")


def _compact_text(value: str, *, limit: int) -> str:
    return " ".join(value.split())[:limit]


def _fallback_tag(content: str) -> str:
    lowered = content.lower()
    if any(word in lowered for word in ("career", "work", "job", "role", "promotion", "mentor")):
        return "career"
    if any(word in lowered for word in ("health", "sleep", "exercise", "doctor")):
        return "health"
    if any(word in lowered for word in ("goal", "plan", "milestone", "project")):
        return "goals"
    return "uncategorized"


def _fallback_summary(summary: str) -> str:
    compact = _compact_text(summary, limit=240)
    return compact or "This document was added to Orbit for future context."


def _auto_weave_connected_buckets(life_item_id: UUID) -> None:
    from app.lifecycle.story_weave import StoryWeaveError, weave_story_bucket

    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT DISTINCT target_id::uuid AS bucket_id
                FROM item_connections
                WHERE source_life_item_id = %s AND target_type = 'story_bucket'
                """,
                (life_item_id,),
            )
            bucket_ids = [row["bucket_id"] for row in cur.fetchall()]

    for bucket_id in bucket_ids:
        try:
            weave_story_bucket(bucket_id)
        except StoryWeaveError:
            continue


def _set_document_bucket_update_text(life_item_id: UUID, connection_summary: str) -> None:
    if not connection_summary.strip():
        return
    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                UPDATE bucket_updates
                SET update_text = %s,
                    updated_at = now()
                WHERE life_item_id = %s
                    AND status = 'pending'
                """,
                (connection_summary, life_item_id),
            )


def _generate_available_unique_name(original_name: str, content: str) -> str:
    base = _generate_unique_name(original_name, content)
    with transaction() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT unique_name
                FROM document_items
                WHERE unique_name = %s OR unique_name LIKE %s
                """,
                (base, f"{base}_%"),
            )
            existing = {row["unique_name"] for row in cur.fetchall()}

    if base not in existing:
        return base

    suffix = 2
    while f"{base}_{suffix}" in existing:
        suffix += 1
    return f"{base}_{suffix}"


def _generate_unique_name(original_name: str, content: str) -> str:
    base = original_name.rsplit(".", 1)[0]
    slug = _slugify(base)
    if slug:
        return slug[:64]
    return f"untitled_doc_{hashlib.sha1(content.encode('utf-8')).hexdigest()[:8]}"


def _slugify(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")
    normalized = re.sub(r"_+", "_", normalized)
    if not normalized or not normalized[0].isalpha():
        return ""
    return normalized


def _extract_upload_text(original_name: str, data: bytes) -> str:
    if not data:
        raise DocumentParseError("The uploaded file is empty.")

    extension = Path(original_name).suffix.lower()
    if extension in {".md", ".markdown", ".txt", ".csv", ".tsv", ".json", ".jsonl", ".xml", ".rtf", ".eml"}:
        text = _decode_text(data)
    elif extension in {".html", ".htm"}:
        text = _strip_html(_decode_text(data))
    elif extension == ".pdf":
        text = _extract_pdf_text(data)
    elif extension == ".docx":
        text = _extract_docx_text(data)
    elif extension == ".doc":
        raise DocumentParseError("Legacy .doc files are not supported yet. Save the file as .docx, PDF, or Markdown.")
    else:
        raise DocumentParseError("Unsupported document type. Upload Markdown, text, PDF, DOCX, CSV, HTML, XML, RTF, or EML.")

    compact = text.strip()
    if not compact:
        raise DocumentParseError("No readable text could be extracted from this file.")
    return compact


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("Unable to decode this text file.")


def _strip_html(content: str) -> str:
    without_scripts = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", content, flags=re.IGNORECASE | re.DOTALL)
    without_tags = re.sub(r"<[^>]+>", " ", without_scripts)
    return re.sub(r"\s+", " ", without_tags).strip()


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard for fresh environments.
        raise DocumentParseError("PDF upload support requires the pypdf package.") from exc

    try:
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # pragma: no cover - parser-specific exceptions vary by PDF.
        raise DocumentParseError("Unable to read text from this PDF.") from exc


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard for fresh environments.
        raise DocumentParseError("DOCX upload support requires the python-docx package.") from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:  # pragma: no cover - parser-specific exceptions vary by file.
        raise DocumentParseError("Unable to read text from this DOCX file.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def _guess_mime_type(original_name: str) -> str:
    extension = Path(original_name).suffix.lower()
    return {
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".txt": "text/plain",
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".csv": "text/csv",
        ".tsv": "text/tab-separated-values",
        ".html": "text/html",
        ".htm": "text/html",
        ".xml": "application/xml",
        ".json": "application/json",
        ".jsonl": "application/jsonl",
        ".rtf": "application/rtf",
        ".eml": "message/rfc822",
    }.get(extension, "application/octet-stream")


def _summary(content: str) -> str:
    compact = " ".join(content.split())
    if len(compact) <= 240:
        return compact
    return f"{compact[:237].rstrip()}..."


def _chunk_text(content: str, chunk_size: int = 1200) -> list[str]:
    compact = content.strip()
    if len(compact) <= chunk_size:
        return [compact]

    chunks: list[str] = []
    start = 0
    while start < len(compact):
        chunks.append(compact[start : start + chunk_size].strip())
        start += chunk_size
    return [chunk for chunk in chunks if chunk]
